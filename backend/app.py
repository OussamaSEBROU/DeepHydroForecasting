# backend/app.py
"""
Single-file Flask backend for Deep Hydro.
- Place pretrained model as backend/standard_model.h5
- Set GOOGLE_API_KEY in environment for Gemini access on Render.
- Endpoints:
  POST /api/upload-data        (form-data file: csv/xlsx)
  GET  /api/data-preview       (?path=...)
  GET  /api/data-stats         (?path=...)
  POST /api/forecast           (JSON: {path, features, time_step, horizon})
  POST /api/generate-report    (JSON: {path, features, time_step, horizon})
  POST /api/ask-ai             (JSON: {query})
  POST /api/chat               (JSON: {message, conversation_id})
  GET  /                         serves frontend index.html (static)
"""

import os
import io
import json
import uuid
import logging
from datetime import datetime
from flask import Flask, request, jsonify, send_file, abort
from flask_cors import CORS
from werkzeug.utils import secure_filename

import numpy as np
import pandas as pd
from sklearn.preprocessing import MinMaxScaler
from tensorflow.keras.models import load_model

from fpdf import FPDF
from docx import Document

# Optional Google Generative AI package. If not available, fallback to error.
try:
    import google.generativeai as genai
    GENAI_AVAILABLE = True
except Exception:
    GENAI_AVAILABLE = False

# Config
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
MODEL_PATH = os.path.join(BASE_DIR, "standard_model.h5")
ALLOWED_EXT = {"csv", "xlsx", "xls"}
MAX_CONTENT_LENGTH = 200 * 1024 * 1024  # 200 MB

# Flask app
app = Flask(__name__, static_folder=os.path.join(BASE_DIR, "..", "frontend"), static_url_path="/")
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH

logging.basicConfig(level=logging.INFO)

# Load model once (pretrained)
_model = None
_scaler_cache = {}  # optional cache if you want to keep scalers per dataset

def load_pretrained_model():
    global _model
    if _model is None:
        if not os.path.exists(MODEL_PATH):
            logging.warning("Pretrained model not found at %s", MODEL_PATH)
            return None
        _model = load_model(MODEL_PATH)
        logging.info("Pretrained model loaded.")
    return _model

# Utility helpers
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def save_uploaded_file(f):
    filename = secure_filename(f.filename)
    uid = uuid.uuid4().hex[:8]
    filename = f"{datetime.utcnow().strftime('%Y%m%d')}_{uid}_{filename}"
    path = os.path.join(UPLOAD_DIR, filename)
    f.save(path)
    return path

def read_table(path, nrows=200):
    ext = path.rsplit(".",1)[1].lower()
    if ext in ("xls","xlsx"):
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path)
    return df

def create_sequences_from_array(arr, time_step):
    X = []
    for i in range(len(arr) - time_step):
        X.append(arr[i:i+time_step])
    return np.array(X)

def forecast_from_model(model, recent_window, horizon, feature_count, scaler=None, feature_columns=None):
    """
    recent_window: array shape (time_step, feature_count)
    horizon: int steps to predict
    scaler: MinMaxScaler used for inverse transform (expected to have same num features + target logic)
    Returns: list of predicted values (original scale if scaler provided; otherwise scaled)
    """
    results = []
    window = recent_window.copy()
    for h in range(horizon):
        x = window.reshape(1, window.shape[0], window.shape[1])
        pred_scaled = model.predict(x, verbose=0).flatten()  # shape (1,)
        # Build next row: we append predicted value as last column or as first depending on original layout.
        # Here we assume target is the first column (commonly Water Level). We'll append predicted value in same feature place.
        next_row = np.zeros(feature_count)
        next_row[0] = pred_scaled[0]  # put predicted value in index 0
        results.append(pred_scaled[0])
        # shift window
        window = np.vstack([window[1:], next_row])
    # if scaler provided, we attempt to inverse transform results by constructing dummy rows
    if scaler is not None and feature_columns:
        # scaler was fitted on columns feature_columns (ordered). Build dummy arrays with same shape.
        inv_results = []
        for val in results:
            dummy = np.zeros((1, len(feature_columns)))
            dummy[0,0] = val  # predicted placed in first column
            try:
                inv = scaler.inverse_transform(dummy)[0,0]
            except Exception:
                inv = val
            inv_results.append(float(inv))
        return inv_results
    return [float(x) for x in results]


# == Endpoints ==

@app.route("/", methods=["GET"])
def index():
    # serve frontend build index.html
    frontend_index = os.path.join(app.static_folder, "index.html")
    if os.path.exists(frontend_index):
        return send_file(frontend_index)
    return "Deep Hydro Backend. Frontend not found.", 200

@app.route("/api/upload-data", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "file missing"}), 400
    f = request.files["file"]
    if f.filename == "":
        return jsonify({"error": "no filename"}), 400
    if not allowed_file(f.filename):
        return jsonify({"error": "file type not allowed"}), 400
    path = save_uploaded_file(f)
    try:
        df = read_table(path)
        # Minimal check : ensure at least one numeric column
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        return jsonify({"message": "uploaded", "path": path, "columns": df.columns.tolist(), "numeric_columns": numeric_cols}), 200
    except Exception as e:
        logging.exception("read_table error")
        return jsonify({"error": str(e)}), 500

@app.route("/api/data-preview", methods=["GET"])
def api_preview():
    path = request.args.get("path")
    if not path or not os.path.exists(path):
        return jsonify({"error": "path missing or not exists"}), 400
    n = int(request.args.get("n", 50))
    try:
        df = read_table(path)
        return jsonify({"columns": df.columns.tolist(), "preview": df.head(n).to_dict(orient="records")})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/data-stats", methods=["GET"])
def api_stats():
    path = request.args.get("path")
    if not path or not os.path.exists(path):
        return jsonify({"error":"path missing or not exists"}), 400
    try:
        df = read_table(path)
        desc = df.describe(include="all").to_dict()
        # extra: correlation matrix for numeric columns
        corr = df.select_dtypes(include=[np.number]).corr().to_dict()
        return jsonify({"describe": desc, "correlation": corr})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/forecast", methods=["POST"])
def api_forecast():
    """
    JSON body:
    {
      "path": "backend/uploads/....csv",
      "feature_columns": ["Water Level","Rainfall",...],   # order must match what model expects
      "time_step": 10,
      "horizon": 7
    }
    """
    payload = request.get_json() or {}
    path = payload.get("path")
    feature_columns = payload.get("feature_columns")
    time_step = int(payload.get("time_step", 10))
    horizon = int(payload.get("horizon", 7))

    if not path or not os.path.exists(path):
        return jsonify({"error":"path missing or file not found"}), 400
    if not feature_columns:
        return jsonify({"error":"feature_columns required"}), 400

    try:
        df = read_table(path)
        # ensure columns exist
        for c in feature_columns:
            if c not in df.columns:
                return jsonify({"error": f"column {c} not found"}), 400

        data_arr = df[feature_columns].values.astype(float)
        # fit scaler on the provided data (MinMax as in original)
        scaler = MinMaxScaler(feature_range=(0,1))
        scaled = scaler.fit_transform(data_arr)

        if len(scaled) < time_step:
            return jsonify({"error": "not enough rows for the given time_step"}), 400

        recent_window = scaled[-time_step:, :]  # shape (time_step, features)
        model = load_pretrained_model()
        if model is None:
            return jsonify({"error": "pretrained model not loaded"}), 500

        preds_scaled = forecast_from_model(model, recent_window, horizon, feature_count=scaled.shape[1], scaler=scaler, feature_columns=feature_columns)
        # timestamps: continue index if Date exists
        timestamps = None
        if "Date" in df.columns or hasattr(df.index, "tolist"):
            try:
                if "Date" in df.columns:
                    last_date = pd.to_datetime(df["Date"].iloc[-1])
                else:
                    last_date = pd.to_datetime(df.index[-1])
                timestamps = [(last_date + pd.Timedelta(days=i+1)).strftime("%Y-%m-%d") for i in range(horizon)]
            except Exception:
                timestamps = None

        return jsonify({"predictions": preds_scaled, "timestamps": timestamps, "horizon": horizon})
    except Exception as e:
        logging.exception("forecast error")
        return jsonify({"error": str(e)}), 500

@app.route("/api/generate-report", methods=["POST"])
def api_generate_report():
    """
    Generate both PDF and Word (docx) report and return as files (zip-like multi-call).
    Body:
      JSON { path, feature_columns, time_step, horizon }
    Returns:
      JSON with base64pdf and base64docx (so frontend can download)
    """
    payload = request.get_json() or {}
    path = payload.get("path")
    feature_columns = payload.get("feature_columns")
    time_step = int(payload.get("time_step", 10))
    horizon = int(payload.get("horizon", 7))

    if not path or not os.path.exists(path):
        return jsonify({"error":"path missing or file not found"}), 400
    if not feature_columns:
        return jsonify({"error":"feature_columns required"}), 400

    try:
        # gather stats
        df = read_table(path)
        desc = df[feature_columns].describe().to_string()
        corr = df[feature_columns].corr().to_string()

        # produce forecast
        forecast_res = api_forecast_json(path, feature_columns, time_step, horizon)
        preds = forecast_res.get("predictions", [])

        # Create PDF
        pdf_bytes = build_pdf_bytes(path, feature_columns, desc, corr, preds)
        # Create Word doc
        docx_bytes = build_docx_bytes(path, feature_columns, desc, corr, preds)

        import base64
        return jsonify({
            "pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8"),
            "docx_base64": base64.b64encode(docx_bytes).decode("utf-8")
        })
    except Exception as e:
        logging.exception("report error")
        return jsonify({"error": str(e)}), 500

def api_forecast_json(path, feature_columns, time_step, horizon):
    # helper to call forecast logic synchronously
    df = read_table(path)
    data_arr = df[feature_columns].values.astype(float)
    scaler = MinMaxScaler(feature_range=(0,1))
    scaled = scaler.fit_transform(data_arr)
    recent_window = scaled[-time_step:, :]
    model = load_pretrained_model()
    preds_scaled = forecast_from_model(model, recent_window, horizon, feature_count=scaled.shape[1], scaler=scaler, feature_columns=feature_columns)
    return {"predictions": preds_scaled}

def build_pdf_bytes(path, feature_columns, desc_text, corr_text, preds):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Deep Hydro — Forecast Report", ln=True, align="C")
    pdf.ln(6)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 6, f"Dataset: {os.path.basename(path)}")
    pdf.ln(4)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 6, "Descriptive Statistics:", ln=True)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 5, desc_text)
    pdf.ln(4)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 6, "Correlation Matrix:", ln=True)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 5, corr_text)
    pdf.ln(6)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 6, "Forecast (next steps):", ln=True)
    pdf.set_font("Arial", "", 10)
    for i, p in enumerate(preds, 1):
        pdf.cell(0, 6, f"Step {i}: {p}", ln=True)
    return pdf.output(dest="S").encode("latin1")

def build_docx_bytes(path, feature_columns, desc_text, corr_text, preds):
    doc = Document()
    doc.add_heading("Deep Hydro — Forecast Report", level=1)
    doc.add_paragraph(f"Dataset: {os.path.basename(path)}")
    doc.add_heading("Descriptive statistics", level=2)
    doc.add_paragraph(desc_text)
    doc.add_heading("Correlation matrix", level=2)
    doc.add_paragraph(corr_text)
    doc.add_heading("Forecast (next steps)", level=2)
    for i, p in enumerate(preds, 1):
        doc.add_paragraph(f"Step {i}: {p}")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

@app.route("/api/ask-ai", methods=["POST"])
def api_ask_ai():
    """
    Body: { "query": "explain ..." }
    Returns: JSON { response: "..." }
    """
    payload = request.get_json() or {}
    query = payload.get("query", "")
    if not query:
        return jsonify({"error":"query required"}), 400

    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        return jsonify({"error":"GOOGLE_API_KEY not set in environment"}), 500

    if not GENAI_AVAILABLE:
        return jsonify({"error":"google.generativeai package not available on server"}), 500

    try:
        genai.configure(api_key=api_key)
        model_ai = genai.GenerativeModel("gemini-pro")
        resp = model_ai.generate_content(query)
        text = getattr(resp, "text", str(resp))
        return jsonify({"response": text})
    except Exception as e:
        logging.exception("genai error")
        return jsonify({"error": str(e)}), 500

# Simple chat endpoint storing ephemeral conversation in memory (for demo)
_conversations = {}

@app.route("/api/chat", methods=["POST"])
def api_chat():
    """
    Body: { "conversation_id": optional, "message": "..." }
    Returns: { conversation_id, reply }
    """
    payload = request.get_json() or {}
    msg = payload.get("message", "")
    if not msg:
        return jsonify({"error":"message required"}), 400
    conv_id = payload.get("conversation_id") or uuid.uuid4().hex
    # Append to conversation
    _conversations.setdefault(conv_id, []).append({"role":"user","text":msg, "ts": datetime.utcnow().isoformat()})
    # Use Gemini if available
    api_key = os.environ.get("GOOGLE_API_KEY")
    if api_key and GENAI_AVAILABLE:
        try:
            genai.configure(api_key=api_key)
            model_ai = genai.GenerativeModel("gemini-pro")
            resp = model_ai.generate_content(msg)
            text = getattr(resp, "text", str(resp))
        except Exception as e:
            text = f"(AI error) {str(e)}"
    else:
        # fallback: echo + simple analysis hint
        text = f"Echo: {msg[:200]}. To enable intelligent replies set GOOGLE_API_KEY in environment and ensure google.generativeai library is installed."
    _conversations[conv_id].append({"role":"assistant","text":text, "ts": datetime.utcnow().isoformat()})
    return jsonify({"conversation_id": conv_id, "reply": text})

# Health
@app.route("/api/health", methods=["GET"])
def health():
    model_present = os.path.exists(MODEL_PATH)
    return jsonify({"status":"ok", "model_present": model_present, "genai_available": GENAI_AVAILABLE})

if __name__ == "__main__":
    # ensure model is attempted to be loaded at start
    load_pretrained_model()
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8000)), debug=False)